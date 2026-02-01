"""RAG (检索增强生成) 引擎"""

from dataclasses import dataclass
from pathlib import Path
from typing import Any, Optional
import hashlib

from auto.core.knowledge.vector_store import (
    VectorStore,
    Document,
    SearchResult,
    get_vector_store,
)


@dataclass
class RAGContext:
    """RAG 上下文"""
    query: str
    documents: list[Document]
    scores: list[float]
    
    def format_context(self, max_length: int = 4000) -> str:
        """格式化为上下文字符串"""
        parts = []
        total_length = 0
        
        for doc, score in zip(self.documents, self.scores):
            content = doc.content.strip()
            
            if total_length + len(content) > max_length:
                # 截断
                remaining = max_length - total_length
                if remaining > 100:
                    content = content[:remaining] + "..."
                else:
                    break
            
            source = doc.metadata.get("source", "unknown")
            parts.append(f"[来源: {source}]\n{content}")
            total_length += len(content) + 50
        
        return "\n\n---\n\n".join(parts)


class RAGEngine:
    """RAG 引擎
    
    负责:
    - 文档索引和管理
    - 语义检索
    - 上下文构建
    """
    
    def __init__(
        self,
        vector_store: Optional[VectorStore] = None,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
    ):
        self._vector_store = vector_store or get_vector_store()
        self._chunk_size = chunk_size
        self._chunk_overlap = chunk_overlap
    
    def _split_text(self, text: str) -> list[str]:
        """分割文本为块"""
        chunks = []
        
        # 按段落分割
        paragraphs = text.split("\n\n")
        
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            if len(current_chunk) + len(para) <= self._chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                # 如果段落本身太长，进一步分割
                if len(para) > self._chunk_size:
                    words = para.split()
                    current_chunk = ""
                    
                    for word in words:
                        if len(current_chunk) + len(word) <= self._chunk_size:
                            current_chunk += word + " "
                        else:
                            if current_chunk:
                                chunks.append(current_chunk.strip())
                            current_chunk = word + " "
                else:
                    current_chunk = para + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    async def index_document(
        self,
        content: str,
        source: str,
        collection: str = "default",
        metadata: Optional[dict] = None,
    ) -> list[str]:
        """索引文档
        
        Args:
            content: 文档内容
            source: 来源标识
            collection: 集合名称
            metadata: 额外元数据
        
        Returns:
            list[str]: 文档块 ID 列表
        """
        # 分割文本
        chunks = self._split_text(content)
        
        # 创建文档
        documents = []
        base_metadata = metadata or {}
        base_metadata["source"] = source
        
        for i, chunk in enumerate(chunks):
            doc_id = hashlib.md5(f"{source}:{i}:{chunk[:50]}".encode()).hexdigest()
            
            doc = Document(
                id=doc_id,
                content=chunk,
                metadata={
                    **base_metadata,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                },
            )
            documents.append(doc)
        
        # 添加到向量存储
        ids = await self._vector_store.add_documents(documents, collection)
        
        return ids
    
    async def index_file(
        self,
        file_path: Path | str,
        collection: str = "default",
        metadata: Optional[dict] = None,
    ) -> list[str]:
        """索引文件
        
        支持: .txt, .md, .pdf (需要额外库)
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 读取文件内容
        content = ""
        
        if path.suffix.lower() == ".pdf":
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(path)
                for page in doc:
                    content += page.get_text()
                doc.close()
            except ImportError:
                raise ImportError("需要安装 PyMuPDF: pip install pymupdf")
        
        elif path.suffix.lower() in [".txt", ".md", ".markdown"]:
            content = path.read_text(encoding="utf-8")
        
        elif path.suffix.lower() in [".doc", ".docx"]:
            try:
                import docx
                doc = docx.Document(path)
                content = "\n".join([p.text for p in doc.paragraphs])
            except ImportError:
                raise ImportError("需要安装 python-docx: pip install python-docx")
        
        else:
            # 尝试作为纯文本读取
            content = path.read_text(encoding="utf-8", errors="ignore")
        
        # 索引内容
        file_metadata = metadata or {}
        file_metadata["file_name"] = path.name
        file_metadata["file_path"] = str(path)
        file_metadata["file_type"] = path.suffix
        
        return await self.index_document(
            content=content,
            source=str(path),
            collection=collection,
            metadata=file_metadata,
        )
    
    async def search(
        self,
        query: str,
        collection: str = "default",
        top_k: int = 5,
        min_score: float = 0.3,
        filter: Optional[dict] = None,
    ) -> RAGContext:
        """搜索相关文档
        
        Args:
            query: 查询文本
            collection: 集合名称
            top_k: 返回数量
            min_score: 最低分数阈值
            filter: 元数据过滤
        
        Returns:
            RAGContext: 检索上下文
        """
        results = await self._vector_store.search(
            query=query,
            collection=collection,
            top_k=top_k,
            filter=filter,
        )
        
        # 过滤低分结果
        filtered_results = [r for r in results if r.score >= min_score]
        
        return RAGContext(
            query=query,
            documents=[r.document for r in filtered_results],
            scores=[r.score for r in filtered_results],
        )
    
    async def query_with_context(
        self,
        query: str,
        collection: str = "default",
        top_k: int = 5,
    ) -> tuple[str, RAGContext]:
        """查询并构建带上下文的提示词
        
        Returns:
            tuple[str, RAGContext]: (构建的提示词, RAG 上下文)
        """
        context = await self.search(query, collection, top_k)
        
        if not context.documents:
            prompt = query
        else:
            context_text = context.format_context()
            prompt = f"""请基于以下参考资料回答问题。

## 参考资料

{context_text}

## 问题

{query}

## 回答要求

1. 优先使用参考资料中的信息
2. 如果参考资料不足以回答，请说明
3. 引用时注明来源"""
        
        return prompt, context
    
    async def delete_by_source(
        self,
        source: str,
        collection: str = "default",
    ) -> int:
        """删除指定来源的文档"""
        # ChromaDB 需要先搜索再删除
        # 这里使用简单实现
        results = await self._vector_store.search(
            query="",
            collection=collection,
            top_k=1000,
            filter={"source": source},
        )
        
        if not results:
            return 0
        
        ids = [r.document.id for r in results]
        return await self._vector_store.delete(ids, collection)
    
    async def list_sources(self, collection: str = "default") -> list[str]:
        """列出所有来源"""
        # 需要遍历所有文档获取来源
        # ChromaDB 不直接支持聚合，这里使用简单实现
        sources = set()
        
        # 搜索所有文档
        results = await self._vector_store.search(
            query="",
            collection=collection,
            top_k=1000,
        )
        
        for r in results:
            source = r.document.metadata.get("source")
            if source:
                sources.add(source)
        
        return list(sources)
    
    async def get_stats(self, collection: str = "default") -> dict:
        """获取统计信息"""
        if hasattr(self._vector_store, "count"):
            count = await self._vector_store.count(collection)
        else:
            results = await self._vector_store.search("", collection, top_k=10000)
            count = len(results)
        
        sources = await self.list_sources(collection)
        
        return {
            "collection": collection,
            "document_count": count,
            "source_count": len(sources),
            "sources": sources[:20],  # 限制返回数量
        }


# 全局 RAG 引擎实例
_rag_engine: Optional[RAGEngine] = None


def get_rag_engine() -> RAGEngine:
    """获取全局 RAG 引擎"""
    global _rag_engine
    if _rag_engine is None:
        _rag_engine = RAGEngine()
    return _rag_engine
